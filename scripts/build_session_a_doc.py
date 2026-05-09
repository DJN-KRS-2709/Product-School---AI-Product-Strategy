#!/usr/bin/env python3
"""Build a formatted Word document for Session A: The Value & Moat facilitation plan."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

SPOTIFY_GREEN = "1DB954"
DARK_HEADER = "4A4A4A"
BORDER_COLOR = "CCCCCC"
AMBER_HEX = "D48B06"
RED_HEX = "C13C37"
BLUE_HEX = "296EB4"

DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MID_GRAY = RGBColor(0x6B, 0x6B, 0x6B)
SOFT_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# -- Base styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lvl in range(1, 4):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.color.rgb = DARK_TEXT
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if lvl == 1 else 14 if lvl == 2 else 10)
    h.paragraph_format.space_after = Pt(6)
    h.font.size = Pt(24 if lvl == 1 else 16 if lvl == 2 else 12)

section = doc.sections[0]
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, m, Cm(2.5))


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))


def set_cell_border(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge in ("top", "bottom", "start", "end"):
        borders.append(parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        ))
    tcPr.append(borders)


def styled_table(headers, rows, col_widths=None, header_color=SPOTIFY_GREEN):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = SOFT_WHITE
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_TEXT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            set_cell_shading(cell, "FFFFFF" if r_idx % 2 == 0 else "F7F7F7")
            set_cell_border(cell, BORDER_COLOR)

    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)

    doc.add_paragraph()
    return table


def callout(text, label="", color_hex=SPOTIFY_GREEN):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    ))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9F9F9"/>'))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    if label:
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(color_hex)
        run.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEXT
    run.italic = True


def divider():
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'</w:pBdr>'
    ))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = "Calibri"
    p.add_run(normal_text).font.name = "Calibri"
    return p


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


Q = chr(0x201C)  # left double curly quote
Qe = chr(0x201D)  # right double curly quote

# ============================================================
# COVER PAGE
# ============================================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SESSION A")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = MID_GRAY
run.font.name = "Calibri"
run.font.letter_spacing = Pt(3)

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run("The Value & Moat")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = DARK_TEXT
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Facilitation Plan \u2014 London Offsite")
run.font.size = Pt(14)
run.font.color.rgb = MID_GRAY
run.font.name = "Calibri"

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Session Leads: Nate Mehari  \u00b7  Dejan Krstic  \u00b7  David Canning",
    "April 22, 2026  \u00b7  09:30 \u2013 10:45  \u00b7  75 minutes",
    "",
    "CONFIDENTIAL \u00b7 FOR INTERNAL USE ONLY",
]:
    run = meta.add_run(line + "\n")
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY
    run.font.name = "Calibri"
    if "CONFIDENTIAL" in line:
        run.font.size = Pt(8)
        run.font.letter_spacing = Pt(2)

doc.add_page_break()

# ============================================================
# CORE QUESTION
# ============================================================

doc.add_heading("Core Question", level=1)

callout(
    f'{Q}What does FinE exist to do, if not write code? Where is our competitive moat '
    f'in a world of abundant code generation?{Qe}'
)

# ============================================================
# INTENT
# ============================================================

doc.add_heading("Intent", level=2)

doc.add_paragraph(
    "Session A is the intellectual spine of Day 2. The Overnight Synthesis gives the room "
    "a shared picture of what surfaced on Day 1 \u2014 the framework mapping, the risk harvest, "
    "the tensions. This session takes that and asks the hardest question: so what are we actually for?"
)

doc.add_paragraph(
    "We are not here to produce an answer. We are here to surface the two or three candidate "
    "answers that this group is willing to defend, stress-test them against each other, and leave "
    "with enough clarity to draft beliefs in Session C."
)

bold_para(
    "The emotional arc: ",
    "start with what we already know (the value chain \u2014 concrete, grounding), move into "
    "what we believe (the moat \u2014 contested, exposing), and land on what we might become "
    "(the new mandate \u2014 forward-looking, uncomfortable). End with a written commitment "
    "so nobody leaves without a personal stake."
)

styled_table(
    ["Follows", "Feeds Into"],
    [["Overnight Synthesis (09:00 \u2013 09:30)",
      "Session B (People, Roles & Culture), Session C (Beliefs Doc)"]],
    col_widths=[3.0, 3.5], header_color=DARK_HEADER,
)

divider()

# ============================================================
# PRE-SESSION CONTEXT
# ============================================================

doc.add_heading("Pre-Session Context", level=1)

doc.add_paragraph(
    "The Alliance Strategy v5 and the March 17 session data surface three distinct moat "
    "arguments. These are not new \u2014 they are already embedded in how FinE talks about itself. "
    "The session\u2019s job is to make them explicit, name which ones are durable, and force a "
    "choice about where we invest."
)

doc.add_heading("The Complexity Facts", level=2)

doc.add_paragraph(
    "These numbers are the floor of the argument \u2014 what makes FinE\u2019s problem space non-trivial. "
    "They are not the moat. They are the reason the moat exists."
)

styled_table(
    ["Dimension", "Scale"],
    [
        ["Annual revenue", "\u20ac15.6B+"],
        ["Annual payouts", "\u20ac11B+"],
        ["Markets", "184"],
        ["Currencies", "98"],
        ["Legal entities", "27"],
        ["Tax returns / year", "2,500+ across 245 markets"],
        ["Tax determinations / year", "2B+"],
        ["Licensors", "800+"],
        ["Creators", "10,000+"],
        ["Revenue calculation methods", "15+"],
        ["Configurable levers", "20+"],
        ["Integrated systems", "100+ (SaaS & in-house)"],
        ["PSPs", "16"],
        ["Settlement files", "177K+"],
        ["Financial transactions / year", "2.4B"],
        ["BETs that touch FinE", "~40%"],
    ],
    col_widths=[2.8, 3.2],
)

doc.add_heading("What the Alliance Strategy Already Claims", level=2)

callout(
    f'{Q}Our organization, FinE, possesses the unique and invaluable combination of deep financial '
    f'domain expertise and a non-negotiable understanding of the trust, compliance, and audit '
    f'requirements of a global finance operation. Therefore, our strategy is not to simply adopt '
    f'GenAI, but to serve as the exclusive, authoritative, and trusted partner in applying this '
    f'technology to solve our highest-impact financial problems.{Qe}',
    label="AI COMPATIBILITY PILLAR", color_hex=BLUE_HEX,
)

callout(
    f'{Q}The deep understanding of why a financial system is structured a certain way, what a '
    f'regulator in a specific jurisdiction actually cares about, what the right question is \u2014 '
    f'that knowledge is FinE\u2019s real value. Are we protecting it?{Qe}',
    label="NAZAR\u00c9 UGLY NOTES", color_hex=BLUE_HEX,
)

callout(
    f'{Q}Financial processes remain fragmented, undocumented, and highly manual, which means AI '
    f'cannot safely accelerate them. If humans already struggle to interpret data, AI will not '
    f'solve the problem.{Qe}',
    label="ALLIANCE STRATEGY DIAGNOSIS", color_hex=BLUE_HEX,
)

doc.add_paragraph(
    "These quotes tell us that FinE already believes its moat is domain expertise + trust + "
    "governance. The question is whether that belief holds up under pressure \u2014 and what it "
    "means for how we organize."
)

divider()

# ============================================================
# THREE STRAW MAN POSITIONS
# ============================================================

doc.add_heading("The Three Straw Man Positions", level=1)

doc.add_paragraph(
    "Rather than presenting one thesis, the session will debate three candidate moat positions. "
    "Each is drawn from the documents. Each implies a different organizational future. The "
    "group\u2019s job is to argue for and against each, then decide which combination they are "
    "willing to stand behind."
)

# --- Position A ---
doc.add_heading("Position A: Domain Expertise", level=2)

bold_para(
    "The claim: ",
    "FinE\u2019s moat is institutional knowledge \u2014 understanding why financial systems are "
    "structured a certain way, what broke in past market events, what regulators actually "
    "enforce, and what the edge cases are in 800+ licensing agreements across 90+ countries."
)

bold_para("Evidence from the strategy:", "")

bullets([
    "Legacy code encodes decisions, not just logic \u2014 rewriting it with AI risks losing "
    "the reasoning behind the structure",
    "15+ revenue calculation methods with 20+ configurable levers \u2014 the permutations "
    "reflect business and regulatory intent, not just technical complexity",
    f'{Q}Can we identify the ten most consequential things FinE knows that are not written '
    f'down anywhere?{Qe} (Nazar\u00e9)',
    f'Publishing royalties domain: {Q}unparalleled level of licensing flexibility{Qe} accepted '
    "as strategic advantage despite complexity cost",
])

callout(
    "Domain expertise lives in people\u2019s heads. It depreciates as people leave. It is the "
    "hardest moat to scale and the easiest to lose through attrition. The Alliance Strategy "
    f'itself identifies this: {Q}As headcount shrinks and AI generates more of the new '
    f'surface\u2026 where does that accumulated knowledge go?{Qe}',
    label="VULNERABILITY", color_hex=RED_HEX,
)

bold_para(
    "Organizational implication: ",
    "FinE as Domain Expert org \u2014 smaller, more senior, invested in knowledge capture and "
    "transfer, measured by depth of understanding rather than throughput."
)

# --- Position B ---
doc.add_heading("Position B: The Trust Layer", level=2)

bold_para(
    "The claim: ",
    "FinE\u2019s moat is the compliance and governance infrastructure \u2014 the ability to ensure "
    "that every financial transaction, in every market, is correct, auditable, and explainable "
    "to a regulator. In a world of AI-generated code, trust and explainability become the "
    "scarce resource."
)

bold_para("Evidence from the strategy:", "")

bullets([
    "SOX compliance, ITGC for all material systems \u2014 non-negotiable, cannot be hand-waved by AI",
    "EU AI Act, BEPS 2.0, e-invoicing mandates across EU and LATAM \u2014 regulatory pressure "
    "is accelerating, not receding",
    f'{Q}Regulators don\u2019t care that an AI wrote it. Can FinE explain every decision in a model '
    f'risk management context?{Qe} (Nazar\u00e9)',
    f'{Q}Compliance is no longer retrofitted after launch. Tax, royalty, and regulatory logic '
    f'are embedded at the point of transaction.{Qe} (Alliance Strategy vision)',
    "The GenAI Safety & Compliance Standard covering 100% of FinE-controlled AI use cases \u2014 already an OKR",
])

callout(
    "Trust is a threshold, not a differentiator. If everyone is compliant, compliance does "
    "not distinguish you. And trust infrastructure is exactly the kind of thing that can be "
    "codified and automated \u2014 potentially by AI itself. The question is whether FinE\u2019s trust "
    "layer is procedural (automatable) or judgmental (requires human interpretation in "
    "ambiguous regulatory contexts).",
    label="VULNERABILITY", color_hex=RED_HEX,
)

bold_para(
    "Organizational implication: ",
    "FinE as Governor org \u2014 setting standards, ensuring safety, maintaining accountability "
    "for systems that others increasingly build. A fundamentally different mandate from builder."
)

# --- Position C ---
doc.add_heading("Position C: System Integration Judgment", level=2)

bold_para(
    "The claim: ",
    "FinE\u2019s moat is the ability to orchestrate 100+ systems (SaaS and in-house) into a "
    "coherent financial operating system. The judgment about how components fit together, "
    "what breaks downstream when upstream changes, and how to absorb business model evolution "
    "without architectural fragility."
)

bold_para("Evidence from the strategy:", "")

bullets([
    f'{Q}FinE operates in one of the most complex technical landscapes at Spotify, integrating '
    f'over 100 technical systems{Qe} \u2014 described as a strategic advantage',
    f'The Platform flywheel: Solutions focus on speed and gross margin ({Q}muscles consuming '
    f'oxygen{Qe}), Platform ensures future efficiency ({Q}restoring oxygen and removing waste '
    f'products{Qe})',
    "~40% of BETs touch FinE \u2014 everything at Spotify eventually flows through the financial "
    "operating system",
    f'The hybrid architecture \u2014 {Q}a deliberate balance of SaaS reach and in-house precision{Qe} '
    "\u2014 requires judgment about when to extend vs. rebuild, when to buy vs. build",
])

callout(
    "Integration knowledge is partly tacit (people know the dependencies) and partly "
    "structural (the architecture itself embodies the judgment). AI is getting better at "
    "understanding system dependencies. If the architecture is well-documented and modular, "
    "the integration judgment may become less scarce. The Alliance Strategy\u2019s own push toward "
    "Golden Paths and modular systems may, over time, commoditize this very moat.",
    label="VULNERABILITY", color_hex=RED_HEX,
)

bold_para(
    "Organizational implication: ",
    "FinE as Platform Architect org \u2014 investing in composable financial infrastructure, "
    "measured by how fast new business models can be financially enabled, focused on the "
    "interfaces between systems rather than the internals."
)

# --- Tension Table ---
doc.add_heading("The Tension Between Positions", level=2)

doc.add_paragraph(
    "These three positions are not mutually exclusive, but they create real resource "
    "allocation tensions:"
)

styled_table(
    ["If we invest heavily in\u2026", "We may underinvest in\u2026", "The risk is\u2026"],
    [
        ["Domain Expertise (A)", "Platform & automation (C)",
         "We become a bottleneck \u2014 the people who know things but can\u2019t scale"],
        ["Trust Layer (B)", "Speed & flexibility (A, C)",
         "We become a gate, not an enabler \u2014 compliance slows the business"],
        ["System Integration (C)", "Knowledge preservation (A)",
         "We build great infrastructure but lose the people who know why it\u2019s built that way"],
    ],
    col_widths=[2.0, 2.0, 2.5], header_color=DARK_HEADER,
)

doc.add_paragraph(
    "The session should surface which combination the group believes in. Most likely it is "
    "some blend \u2014 but the priority matters because it determines what we hire for, what we "
    "measure, and what we protect."
)

divider()

# ============================================================
# SESSION FLOW
# ============================================================

doc.add_heading("Session Flow: 75 Minutes", level=1)
doc.add_page_break()

# --- Part 1 ---
doc.add_heading("Part 1 \u2014 The Value Chain  (20 min)", level=2)

bold_para(
    "Purpose: ",
    "Build a shared, honest picture of where FinE\u2019s value actually lives today \u2014 before "
    "anyone argues about where it should live."
)
bold_para("Format: ", "Silent first, then open.")

styled_table(
    ["Time", "Activity"],
    [
        ["0:00 \u2013 0:02",
         "Frame the session. Connect to the Overnight Synthesis: the mapping told us where "
         "AI hits hardest and where we need to protect. This session asks a different question "
         "\u2014 not what is changing, but what are we for. Read the core question aloud."],
        ["0:02 \u2013 0:08",
         "Silent stickies. Each person writes 2 things FinE is uniquely valuable at \u2014 not "
         "things we do, but things nobody else could do in our absence. Physical cards or "
         "digital. No talking until all cards are posted."],
        ["0:08 \u2013 0:15",
         "Cluster and name. Facilitator groups stickies on a shared board in real time. Read "
         "them aloud as you cluster. Name each cluster with a short label."],
        ["0:15 \u2013 0:20",
         "React. Full group: what surprises you? What is missing? Capture additions. Flag "
         "the 2\u20133 clusters that generated the most cards \u2014 these are your signal for Part 2."],
    ],
    col_widths=[1.2, 5.3], header_color=SPOTIFY_GREEN,
)

callout(
    "The clusters that emerge will likely map loosely to the three straw man positions. If "
    "they don\u2019t \u2014 that\u2019s more interesting. Name the deviation.",
    label="FACILITATOR NOTE", color_hex=AMBER_HEX,
)

callout(
    "What good looks like: Specific answers (e.g. we understand why the Swedish royalty model "
    "uses a different allocation method than the US one) vs. generic answers (we know finance). "
    "Push back on the generic ones: who else could do that? What makes it ours?",
    label="COACHING", color_hex=AMBER_HEX,
)

# --- Part 2 ---
doc.add_heading("Part 2 \u2014 The Moat  (25 min)", level=2)

bold_para(
    "Purpose: ",
    "Move from what we do to what is defensible. Force the group to argue for and "
    "against specific moat positions, and surface which ones they believe are durable vs. at risk."
)
bold_para("Format: ", "Structured debate, full group.")

styled_table(
    ["Time", "Activity"],
    [
        ["0:20 \u2013 0:25",
         "Present the three positions. Project them on screen. Read each one in 60 seconds "
         "\u2014 the claim and the vulnerability. Don\u2019t argue for any of them. Let the room sit "
         "with the tension."],
        ["0:25 \u2013 0:30",
         "Temperature check. Each person holds up 1, 2, or 3 fingers for which position they "
         "believe is FinE\u2019s primary moat. Count the room. If it\u2019s split \u2014 that\u2019s the data. "
         "If it\u2019s lopsided \u2014 probe the minority view."],
        ["0:30 \u2013 0:40",
         "Open debate on the two most contested positions.\n"
         "\u2022 Is this moat durable in 3 years, or is AI eroding it?\n"
         "\u2022 If we lost this moat tomorrow, what would break first?\n"
         "\u2022 Which of these moats requires people vs. which can be embedded in systems?"],
        ["0:40 \u2013 0:45",
         "Name the clusters. Facilitator captures on the shared board: which moats the room "
         "considers durable, which are at risk, which are contested. These carry forward to Session C."],
    ],
    col_widths=[1.2, 5.3], header_color=SPOTIFY_GREEN,
)

bold_para("Provocations if the conversation stalls:", "")
bullets([
    f'{Q}The ability to write Python is no longer scarce. What is?{Qe}',
    f'{Q}A 10x increase in code output with a 2x increase in subtle errors is a net negative '
    f'in a financial context where correctness is existential.{Qe}',
    f'{Q}Who in this room can explain an AI-generated risk model to a regulator? '
    f'Who can do it in two years?{Qe}',
    f'{Q}If we asked our Finance partners what FinE is uniquely valuable at \u2014 would they say '
    f'the same things we just said?{Qe}',
])

bold_para("Provocations if the conversation goes abstract:", "")
bullets([
    f'{Q}Can you make that concrete? What does this look like for a FinE engineer '
    f'on a Tuesday in October?{Qe}',
    f'{Q}Name one specific system or process where this moat is real, right now.{Qe}',
])

# --- Part 3 ---
doc.add_heading("Part 3 \u2014 The New Mandate  (20 min)", level=2)

bold_para(
    "Purpose: ",
    "Confront the question that sits underneath the moat debate \u2014 is FinE still a builder, "
    "or is it becoming something else?"
)
bold_para("Format: ", "Open debate, facilitator-guided.")

styled_table(
    ["Time", "Activity"],
    [
        ["0:45 \u2013 0:50",
         "Frame the shift. The Alliance Strategy describes FinE as the financial operating "
         "system behind every Spotify launch. But the Nazar\u00e9 provocation asks whether FinE\u2019s "
         "role is shifting from builder to governor. Let\u2019s take that seriously.\n\n"
         "Project: If Finance could build everything themselves tomorrow, what would they "
         "still need FinE for?"],
        ["0:50 \u2013 1:00",
         "Structured responses. Go around the room. Each person answers in one sentence: "
         "Finance would still need FinE for ______.\n"
         "No preamble, no hedging. Facilitator captures verbatim."],
        ["1:00 \u2013 1:05",
         "The builder vs. governor tension. Read the Alliance Strategy\u2019s own framing: "
         "FinE\u2019s role may shift from builder to governor \u2014 setting standards, ensuring "
         "safety, and maintaining accountability for systems that Finance increasingly builds "
         "and operates itself. That is a fundamentally different organization.\n\n"
         "Ask: What do we gain by becoming a governor? What do we lose? Capture both sides."],
    ],
    col_widths=[1.2, 5.3], header_color=SPOTIFY_GREEN,
)

callout(
    f'{Q}Where does FinE end if Finance starts building with MCPs directly?{Qe}',
    label="RESERVE QUESTION", color_hex=AMBER_HEX,
)

callout(
    f'{Q}Are we building toward a smaller, deeper, more accountable org \u2014 or a same-size org '
    f'with dramatically expanded scope? Both are viable. We cannot pursue both at once.{Qe}',
    label="RESERVE QUESTION", color_hex=AMBER_HEX,
)

# --- Synthesis ---
doc.add_heading("Synthesis  (10 min)", level=2)

bold_para(
    "Purpose: ",
    "Lock in a personal commitment from every person in the room. These statements become "
    "raw material for Session C (Beliefs Doc)."
)

styled_table(
    ["Time", "Activity"],
    [
        ["1:05 \u2013 1:10",
         "Individual written completion. Each person writes on a card or in the shared doc:\n\n"
         "FinE\u2019s core value in 2027 is ______.\n\n"
         "No skipping. No hedging. One sentence. The constraint is the point \u2014 it forces a choice."],
        ["1:10 \u2013 1:15",
         "Read aloud. Facilitator collects and reads each one aloud without attribution. "
         "No discussion. Just listen. Note the convergence and the divergence. "
         "Hand both to Session C."],
    ],
    col_widths=[1.2, 5.3], header_color=SPOTIFY_GREEN,
)

callout(
    "Resist the temptation to summarize or synthesize in the moment. Session C does that "
    "work. The value here is the raw, unedited statements. If you smooth them out, you lose "
    "the signal.",
    label="FACILITATOR NOTE", color_hex=AMBER_HEX,
)

divider()

# ============================================================
# MATERIALS & LOGISTICS
# ============================================================

doc.add_heading("Materials & Logistics", level=1)

styled_table(
    ["Item", "Owner", "Status"],
    [
        ["Three straw man positions \u2014 one-pager or projected slide", "Dejan", "To prepare"],
        ["Sticky notes / cards + markers or shared digital board", "Nate",
         "To confirm with offsite logistics"],
        ["Core question projected on screen at entry", "David", "To prepare slide"],
        ["Shared board / canvas for facilitator captures", "Nate", "To set up"],
        ["Printed complexity facts table (optional)", "Dejan", "To prepare"],
    ],
    col_widths=[3.5, 1.0, 2.0], header_color=DARK_HEADER,
)

divider()

# ============================================================
# EXPECTED OUTPUTS
# ============================================================

doc.add_heading("Expected Outputs", level=1)

doc.add_paragraph("What this session must hand off to Session B and Session C:")

styled_table(
    ["Output", "What it is", "Where it goes"],
    [
        ["Named value themes",
         "The 2\u20133 clusters from Part 1 that got the most energy",
         "Session C: Beliefs Doc structure"],
        ["Moat durability assessment",
         "Which moats the room considers durable, at risk, or contested",
         "Session C: What we believe FinE exists to do"],
        ["Builder vs. governor split",
         "The room\u2019s position on FinE\u2019s mandate shift",
         "Session B: Role implications\nSession D: Decisions"],
        ["Core value statements",
         "Individual written completions, unedited",
         "Session C: Raw material for beliefs drafting"],
        ["Unresolved tensions",
         "Anything the group could not agree on",
         "Session D: Named for owner assignment"],
    ],
    col_widths=[2.0, 2.2, 2.3], header_color=DARK_HEADER,
)

divider()

# ============================================================
# COORDINATION NOTES
# ============================================================

doc.add_heading("Coordination Notes for Nate & David", level=1)

bold_para(
    "Upstream dependency: ",
    "This session relies on the Overnight Synthesis having clearly presented the Day 1 "
    "framework mapping and risk harvest. Coordinate with Dylan / Ben / Viktor to ensure the "
    "synthesis names the items in the Invest & Protect quadrant \u2014 those are the starting "
    "point for our value chain discussion."
)

bold_para(
    "Downstream handoff: ",
    "Session B (People, Roles & Culture) opens with: Session A just explored what FinE "
    "exists to do \u2014 our value and our moat. Now we turn to the people side. Coordinate "
    "with Joakim / Anil / Ximena on what we expect to hand them \u2014 specifically whether we "
    "land on builder, governor, or both with tension."
)

doc.add_heading("Proposed Facilitation Split", level=2)

styled_table(
    ["Part", "Facilitator", "Role"],
    [
        ["Part 1 \u2014 Value Chain", "Nate", "Opens the session, clusters the stickies, names themes"],
        ["Part 2 \u2014 The Moat", "Dejan", "Presents the three positions, manages the debate"],
        ["Part 3 \u2014 New Mandate", "David",
         "Frames the builder / governor question, runs structured responses"],
        ["Synthesis", "All three", "Collect cards, read aloud, hand off to Session C"],
    ],
    col_widths=[1.8, 1.0, 3.7], header_color=SPOTIFY_GREEN,
)

doc.add_heading("Prep Timeline", level=2)

styled_table(
    ["By", "What", "Who"],
    [
        ["April 14", "Align on this facilitation plan", "All three"],
        ["April 16", "First draft of the straw man positions slide / one-pager", "Dejan"],
        ["April 18", "Materials confirmed with offsite logistics", "Nate"],
        ["April 20", "Dry run or async walkthrough of the flow", "All three"],
    ],
    col_widths=[1.2, 3.5, 1.8], header_color=DARK_HEADER,
)

divider()

# ============================================================
# PROVOCATION BANK
# ============================================================

doc.add_heading("Appendix: Provocation Bank", level=1)

doc.add_paragraph(
    "Questions drawn from both the Alliance Strategy v5 and the Nazar\u00e9 pre-session data. "
    "Organized by failure mode \u2014 deploy as needed."
)

doc.add_heading("If the room is being too safe or generic", level=2)
bullets([
    f'{Q}What is the most important thing FinE knows that is not written down anywhere?{Qe}',
    f'{Q}If FinE were starting from scratch today, would we build the same organization?{Qe}',
    f'{Q}What is the thing we are most reluctant to say out loud in this room?{Qe}',
])

doc.add_heading("If the room is being too abstract", level=2)
bullets([
    f'{Q}Can you make that concrete? What does this look like for a FinE engineer '
    f'on a Tuesday in October?{Qe}',
    f'{Q}Name the system. Name the process. Name the person.{Qe}',
    f'{Q}What does this mean for the engineer who joins FinE in 6 months?{Qe}',
])

doc.add_heading("If the room collapses into solution mode too early", level=2)
bullets([
    f'{Q}We\u2019re getting into the plan. Let\u2019s park that and stay with the belief: '
    f'do we agree on what we\u2019re optimizing for?{Qe}',
    f'{Q}Before we talk about how \u2014 do we agree on what we are trying to do?{Qe}',
])

doc.add_heading("If anxiety rises", level=2)
bullets([
    f'{Q}It sounds like there\u2019s real fear in the room. That\u2019s right and it belongs here. '
    f'Let\u2019s name what we\u2019re afraid of before we try to solve it.{Qe}',
])

doc.add_heading("Nuclear option", level=2)

callout(
    f'{Q}The productivity gains are real. But we are reducing the number of humans who deeply '
    f'understand our most critical systems, at the same moment those systems are being built '
    f'in new ways we don\u2019t yet fully understand. Are we taking that risk knowingly?{Qe}',
    label="DEPLOY IF STUCK", color_hex=RED_HEX,
)

# ============================================================
# SAVE
# ============================================================

output_path = os.path.expanduser(
    "~/Downloads/Session A - The Value and Moat - Facilitation Plan.docx"
)
doc.save(output_path)
print(f"Saved to: {output_path}")
